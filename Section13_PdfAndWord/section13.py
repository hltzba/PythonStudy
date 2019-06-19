# encoding:utf-8
#! python3
# section13.py - 
import PyPDF2,docx
from docx.enum.style import WD_STYLE_TYPE
def getPDFText():
    fileobj=open('E:\EBook\F. Scott Fitzgerald\The Great Gatsby (4)\\The Great Gatsby - F. Scott Fitzgerald.pdf','rb')
    pdfReader=PyPDF2.PdfFileReader(fileobj)
    print('文档页数:'+str(pdfReader.numPages))
    print('是否加密：'+str(pdfReader.isEncrypted))
    page1obj=pdfReader.getPage(1)
    print(page1obj.extractText())

def getWordText():
    doc=docx.Document('D:\\党员学习\\王志敏三十讲通篇心得.docx')
    print(len(doc.paragraphs))
    '''
    for i in range(len(doc.paragraphs)):
        print(doc.paragraphs[i].text)
        print('-'.center(24,'-'))
    '''
    print(len(doc.paragraphs[2].runs))
    print(doc.paragraphs[2].runs[0].text)
    print(doc.paragraphs[0].style)

def setWordStyle():
    doc=docx.Document()
    doc.add_paragraph('我是标题',style='Heading 1')
    #doc.paragraphs[0].text='我是标题'
    #doc.paragraphs[0].style='Heading1'
    doc.save('style.docx')

def getAllWordStyles():
    styles=docx.Document().styles
    paragraph_styles = []
    for s in styles:
        if s.type==WD_STYLE_TYPE.PARAGRAPH:
            paragraph_styles.append(s)
    for s in paragraph_styles:
        print(s.name)

def addPictureToWord():
    doc=docx.Document()
    doc.add_picture('H:\\BaiduYunDownload\\GamesHeadPhoto\\AssassinsCreed001.png',
        width=docx.shared.Cm(4),height=docx.shared.Cm(4))
    doc.save('picdemo.docx')

#getPDFText()
#getWordText()
#setWordStyle()
#getAllWordStyles()
addPictureToWord()